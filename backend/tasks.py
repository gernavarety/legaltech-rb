"""
Celery задачи обработки договоров.
Основная задача: скачать файл → извлечь текст → найти нормы → Claude API → отчёт DOCX.
"""
import io
import json
import re
import time
from typing import List, Optional
from datetime import datetime

import anthropic
import fitz  # PyMuPDF
import openai
from celery import Celery
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

import database
import storage
from config import get_settings
from models import AnalysisResult, RiskItem

settings = get_settings()

# --- Celery приложение ---
celery_app = Celery(
    "legaltech",
    broker=settings.redis_url,
    backend=settings.redis_url,
)
celery_app.conf.update(
    task_serializer="json",
    accept_content=["json"],
    result_serializer="json",
    timezone="Europe/Minsk",
    enable_utc=True,
    task_track_started=True,
    task_soft_time_limit=300,
    task_time_limit=360,
)

# --- Системный промпт для юриста РБ ---
SYSTEM_PROMPT = """Ты опытный юрист Республики Беларусь со специализацией в гражданском и корпоративном праве.

ПРАВИЛА:
- Всегда ссылайся на конкретные статьи НПА Республики Беларусь
- Используй только белорусское право (ГК РБ, ТК РБ, ХПК РБ и др.)
- Никогда не применяй российское право
- Каждый риск обоснуй ссылкой на статью

ФОРМАТ ОТВЕТА (строго JSON, без markdown-блоков, только чистый JSON):
{
  "contract_type": "тип договора",
  "overall_risk": "низкий|средний|высокий",
  "summary": "краткое резюме 2-3 предложения",
  "risks": [
    {
      "level": "высокий|средний|низкий",
      "clause": "пункт договора или цитата",
      "issue": "описание проблемы",
      "law_reference": "ст. 402 ГК РБ",
      "recommendation": "конкретная рекомендация"
    }
  ],
  "missing_clauses": ["список важных условий которых нет"],
  "needs_lawyer": true
}"""


# ─── Вспомогательные функции ────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Извлекает весь текст из PDF через PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()
    full_text = "\n".join(pages)
    logger.info(f"PDF: извлечено {len(full_text)} символов, {len(pages)} страниц")
    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Извлекает текст из DOCX через python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    full_text = "\n".join(paragraphs)
    logger.info(f"DOCX: извлечено {len(full_text)} символов")
    return full_text


def split_into_chunks(text: str, chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """
    Разбивает текст на чанки по ~chunk_size токенов с перекрытием overlap.
    Используем приблизительный подсчёт: 1 токен ≈ 4 символа.
    """
    char_size = chunk_size * 4
    char_overlap = overlap * 4

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + char_size, text_len)
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk)
        start += char_size - char_overlap

    logger.info(f"Текст разбит на {len(chunks)} чанков")
    return chunks


def create_embeddings(texts: List[str]) -> List[List[float]]:
    """Создаёт эмбеддинги через OpenAI text-embedding-3-small."""
    client = openai.OpenAI(api_key=settings.openai_api_key)
    # OpenAI API принимает до 2048 текстов за раз
    all_embeddings = []
    batch_size = 100

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        response = client.embeddings.create(
            model=settings.embedding_model,
            input=batch,
        )
        batch_embeddings = [item.embedding for item in response.data]
        all_embeddings.extend(batch_embeddings)

    logger.info(f"Создано {len(all_embeddings)} эмбеддингов")
    return all_embeddings


def create_single_embedding(text: str) -> List[float]:
    """Создаёт один эмбеддинг для запроса."""
    client = openai.OpenAI(api_key=settings.openai_api_key)
    response = client.embeddings.create(
        model=settings.embedding_model,
        input=[text[:8000]],  # Ограничиваем длину запроса
    )
    return response.data[0].embedding


def analyze_with_claude(contract_text: str, law_context: str) -> AnalysisResult:
    """
    Отправляет договор в Claude API с контекстом норм РБ.
    Retry логика: 3 попытки с паузами.
    """
    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    user_message = f"""Проанализируй следующий договор на соответствие законодательству Республики Беларусь.

РЕЛЕВАНТНЫЕ НОРМЫ ЗАКОНОДАТЕЛЬСТВА РБ:
{law_context}

ТЕКСТ ДОГОВОРА:
{contract_text[:15000]}

Верни строго JSON без markdown-блоков."""

    last_error = None
    for attempt in range(1, 4):
        try:
            logger.info(f"Claude API: попытка {attempt}/3")
            message = client.messages.create(
                model=settings.claude_model,
                max_tokens=4096,
                timeout=settings.claude_timeout_seconds,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_message}],
            )
            raw_text = message.content[0].text.strip()

            # Убираем возможные markdown-обёртки
            raw_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
            raw_text = re.sub(r"\s*```$", "", raw_text)

            data = json.loads(raw_text)
            result = AnalysisResult(**data)
            logger.info(f"Claude API: успешный ответ, риск={result.overall_risk}")
            return result

        except json.JSONDecodeError as e:
            last_error = e
            logger.warning(f"Claude API попытка {attempt}: ошибка парсинга JSON: {e}")
            time.sleep(5 * attempt)
        except anthropic.APIError as e:
            last_error = e
            logger.warning(f"Claude API попытка {attempt}: API ошибка: {e}")
            time.sleep(10 * attempt)

    raise RuntimeError(f"Claude API недоступен после 3 попыток: {last_error}")


def _set_cell_background(cell, hex_color: str):
    """Устанавливает цвет фона ячейки таблицы DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _risk_color(level: str) -> str:
    """Возвращает HEX-цвет фона по уровню риска."""
    mapping = {
        "высокий": "FFCCCC",
        "средний": "FFF3CD",
        "низкий": "D4EDDA",
    }
    return mapping.get(level.lower(), "FFFFFF")


def generate_docx_report(result: AnalysisResult, filename: str) -> bytes:
    """
    Генерирует DOCX-отчёт из результатов анализа.
    Возвращает байты готового файла.
    """
    doc = Document()

    # --- Стили ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Заголовок ---
    title = doc.add_heading(f"Анализ договора — {result.contract_type}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Дата и общий риск ---
    doc.add_paragraph(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Исходный файл: {filename}")

    risk_para = doc.add_paragraph()
    risk_para.add_run("Общий уровень риска: ").bold = True
    risk_run = risk_para.add_run(result.overall_risk.upper())
    risk_run.bold = True
    colors = {"высокий": RGBColor(0xCC, 0x00, 0x00), "средний": RGBColor(0xFF, 0x99, 0x00), "низкий": RGBColor(0x00, 0x80, 0x00)}
    risk_run.font.color.rgb = colors.get(result.overall_risk.lower(), RGBColor(0, 0, 0))
    risk_run.font.size = Pt(14)

    doc.add_paragraph()

    # --- Резюме ---
    doc.add_heading("Резюме", level=2)
    doc.add_paragraph(result.summary)
    doc.add_paragraph()

    # --- Таблица рисков ---
    doc.add_heading("Выявленные риски", level=2)

    if result.risks:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Заголовок таблицы
        hdr_cells = table.rows[0].cells
        headers = ["Уровень", "Пункт договора", "Проблема", "Норма РБ", "Рекомендация"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            _set_cell_background(hdr_cells[i], "D9D9D9")

        # Строки с рисками
        for risk in result.risks:
            row_cells = table.add_row().cells
            row_cells[0].text = risk.level
            row_cells[1].text = risk.clause[:300]
            row_cells[2].text = risk.issue
            row_cells[3].text = risk.law_reference
            row_cells[4].text = risk.recommendation

            # Цвет фона строки по уровню риска
            bg = _risk_color(risk.level)
            for cell in row_cells:
                _set_cell_background(cell, bg)

        # Ширина колонок
        for i, width in enumerate([1.0, 2.5, 3.0, 1.5, 3.0]):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    else:
        doc.add_paragraph("Существенных рисков не выявлено.")

    doc.add_paragraph()

    # --- Отсутствующие условия ---
    if result.missing_clauses:
        doc.add_heading("Отсутствующие условия", level=2)
        for clause in result.missing_clauses:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clause)
        doc.add_paragraph()

    # --- Вывод ---
    doc.add_heading("Вывод", level=2)
    conclusion = doc.add_paragraph()
    conclusion.add_run("Необходима консультация юриста: ").bold = True
    answer = "ДА" if result.needs_lawyer else "НЕТ"
    answer_run = conclusion.add_run(answer)
    answer_run.bold = True
    answer_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00) if result.needs_lawyer else RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Подготовлено: LexAI.by")
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info("DOCX-отчёт сгенерирован")
    return buffer.read()


# ─── Основная Celery задача ──────────────────────────────────────────────────

import asyncio
from functools import wraps


def run_async(coro):
    """Запускает async функцию из синхронного Celery контекста."""
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


@celery_app.task(bind=True, name="tasks.process_contract", max_retries=0)
def process_contract(self, task_id: str, file_key: str, filename: str):
    """
    Основная задача обработки договора:
    1. Скачать файл из R2
    2. Извлечь текст
    3. Создать эмбеддинги чанков
    4. Найти релевантные нормы РБ
    5. Отправить в Claude API
    6. Сгенерировать DOCX
    7. Загрузить отчёт в R2
    """
    logger.info(f"[{task_id}] Начало обработки файла: {filename}")

    # Статус: в обработке
    run_async(database.update_document_status(task_id, "processing"))

    try:
        # 1. Скачиваем файл из R2
        logger.info(f"[{task_id}] Шаг 1: Скачивание файла из R2")
        file_bytes = storage.download_file(file_key)

        # 2. Извлекаем текст
        logger.info(f"[{task_id}] Шаг 2: Извлечение текста")
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            contract_text = extract_text_from_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            contract_text = extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")

        if len(contract_text.strip()) < 100:
            raise ValueError("Текст договора слишком короткий или не удалось извлечь")

        # 3. Разбиваем на чанки
        logger.info(f"[{task_id}] Шаг 3: Разбивка на чанки")
        chunks = split_into_chunks(
            contract_text,
            chunk_size=settings.chunk_size_tokens,
            overlap=settings.chunk_overlap_tokens,
        )

        # 4. Создаём эмбеддинги для поиска по правовой базе
        logger.info(f"[{task_id}] Шаг 4: Создание эмбеддингов")
        # Используем первые 3000 символов как запрос для поиска норм
        search_text = contract_text[:3000]
        query_embedding = create_single_embedding(search_text)

        # 5. Ищем релевантные нормы РБ в pgvector
        logger.info(f"[{task_id}] Шаг 5: Поиск релевантных норм РБ")
        law_norms = run_async(database.search_similar_laws(query_embedding, top_k=10))

        if law_norms:
            law_context = "\n\n".join([
                f"[{n['document_name']}] {n.get('article_number', '')} {n.get('article_title', '')}\n{n['text'][:500]}"
                for n in law_norms
            ])
            logger.info(f"[{task_id}] Найдено {len(law_norms)} релевантных норм")
        else:
            law_context = "Правовая база РБ временно недоступна. Используй общие знания белорусского права."
            logger.warning(f"[{task_id}] Правовая база пуста, используем общие знания")

        # 6. Анализ через Claude API
        logger.info(f"[{task_id}] Шаг 6: Анализ через Claude API")
        analysis = analyze_with_claude(contract_text, law_context)

        # 7. Генерируем DOCX-отчёт
        logger.info(f"[{task_id}] Шаг 7: Генерация DOCX-отчёта")
        docx_bytes = generate_docx_report(analysis, filename)

        # 8. Загружаем отчёт в R2
        logger.info(f"[{task_id}] Шаг 8: Загрузка отчёта в R2")
        report_key = storage.upload_report(docx_bytes, task_id)

        # 9. Сохраняем результаты в БД
        run_async(database.update_document_status(
            task_id,
            status="done",
            contract_type=analysis.contract_type,
            overall_risk=analysis.overall_risk,
            result_json=analysis.model_dump(),
            report_url=report_key,
        ))

        logger.info(f"[{task_id}] Обработка завершена успешно")
        return {"status": "done", "task_id": task_id}

    except Exception as e:
        error_msg = str(e)
        logger.error(f"[{task_id}] Ошибка обработки: {error_msg}")
        run_async(database.update_document_status(
            task_id,
            status="error",
            error_text=error_msg,
        ))
        raise
