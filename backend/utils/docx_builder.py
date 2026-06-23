"""
Построитель DOCX из текста сгенерированного AI.
Форматирует документ по стандартам делопроизводства РБ:
- Шрифт Times New Roman, 12pt, межстрочный интервал 1.5
- Заголовок по центру жирным
- Нумерованные разделы
- Место для подписей и реквизитов сторон
"""

import re
import os
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, bold: bool = False, size_pt: int = 12):
    """Применяет Times New Roman с нужными параметрами к run."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # CyrillicFont (rFonts cs) — нужен для корректного отображения кириллицы
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:cs"), "Times New Roman")


def _set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before_pt: int = 0, space_after_pt: int = 6,
                           line_spacing: float = 1.5):
    """Выравнивание, отступы и межстрочный интервал."""
    fmt = paragraph.paragraph_format
    fmt.alignment = alignment
    from docx.shared import Pt
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing = line_spacing


def _add_horizontal_line(doc: Document):
    """Добавляет горизонтальную разделительную линию."""
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pb.append(bottom)
    pPr.append(pb)
    return paragraph


def build_docx_from_text(text: str, document_name: str) -> str:
    """
    Принимает текст от Claude и создаёт отформатированный DOCX.
    Возвращает путь к временному файлу (вызывающая сторона удаляет его).
    """
    doc = Document()

    # Поля страницы: верхнее/нижнее 2 см, левое 3 см, правое 1.5 см
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    # Стиль Normal
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    lines = text.strip().split("\n")
    i = 0
    title_done = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Пустая строка
        if not line:
            i += 1
            continue

        # Заголовок документа — первые непустые строки до первого раздела
        if not title_done and not re.match(r"^\d+[\.\)]", line):
            p = doc.add_paragraph()
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.5)
            run = p.add_run(line)
            _set_font(run, bold=True, size_pt=14)
            i += 1
            continue

        # Разделительная линия после заголовка
        if not title_done and re.match(r"^\d+[\.\)]", line):
            title_done = True
            doc.add_paragraph()  # отступ

        # Заголовок раздела вида "1. ОБЩИЕ ПОЛОЖЕНИЯ" или "1. Общие положения"
        section_match = re.match(r"^(\d+)\.\s+([А-ЯA-Z].{0,80})$", line)
        if section_match and len(line) < 100:
            p = doc.add_paragraph()
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 12, 6, 1.5)
            run = p.add_run(line)
            _set_font(run, bold=True, size_pt=12)
            i += 1
            continue

        # Подпункты вида "1.1." или "1.1)"
        subsection_match = re.match(r"^(\d+\.\d+[\.\)])", line)
        if subsection_match:
            p = doc.add_paragraph()
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 1.5)
            p.paragraph_format.first_line_indent = Cm(1.25)
            run = p.add_run(line)
            _set_font(run, bold=False, size_pt=12)
            i += 1
            continue

        # Блок подписей — специальное форматирование таблицей
        if re.search(r"(ПОДПИСИ|РЕКВИЗИТЫ|Подписи сторон|СТОРОНЫ)", line, re.I):
            doc.add_paragraph()
            _add_horizontal_line(doc)
            p = doc.add_paragraph()
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 12, 6)
            run = p.add_run(line)
            _set_font(run, bold=True, size_pt=12)
            i += 1
            # Собираем строки блока подписей в таблицу 2 колонки
            signature_lines = []
            while i < len(lines) and lines[i].strip():
                signature_lines.append(lines[i])
                i += 1
            if signature_lines:
                _add_signature_table(doc, signature_lines)
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 1.5)
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(line)
        _set_font(run, bold=False, size_pt=12)
        i += 1

    # Сохраняем во временный файл
    tmp = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx",
        prefix=f"{_slugify(document_name)}_"
    )
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _add_signature_table(doc: Document, lines: list[str]):
    """
    Создаёт таблицу реквизитов сторон: левый столбец — сторона 1, правый — сторона 2.
    """
    # Делим строки пополам
    mid = len(lines) // 2
    left_lines = lines[:mid]
    right_lines = lines[mid:]

    table = doc.add_table(rows=max(len(left_lines), len(right_lines)), cols=2)
    table.style = "Table Grid"

    # Убираем внешние границы — только пространство
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "nil")
                tc_borders.append(border)
            tc_pr.append(tc_borders)

    for idx, line in enumerate(left_lines):
        if idx < len(table.rows):
            cell = table.cell(idx, 0)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(line)
            _set_font(run, size_pt=11)

    for idx, line in enumerate(right_lines):
        if idx < len(table.rows):
            cell = table.cell(idx, 1)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(line)
            _set_font(run, size_pt=11)


def _slugify(text: str) -> str:
    """Очищает название для использования в имени файла."""
    text = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:50]
