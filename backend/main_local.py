"""
Локальный бэкенд LegalTech RB.
Переключение AI через переменные окружения — без смены кода.

Быстрый старт (Ollama, бесплатно):
    LLM_PROVIDER=ollama python main_local.py

Переключить на Groq (тоже бесплатно):
    LLM_PROVIDER=groq GROQ_API_KEY=gsk_... python main_local.py

Переключить на Claude:
    LLM_PROVIDER=anthropic ANTHROPIC_API_KEY=sk-ant-... python main_local.py
"""
import asyncio
import io
import json
import os
import re
import sqlite3
import tempfile
import uuid
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Optional

import aiosqlite
import fitz
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from loguru import logger

from providers import get_llm_provider, get_embedding_provider

# ── Конфигурация ─────────────────────────────────────────────────
UPLOAD_DIR = Path(tempfile.gettempdir()) / "legaltech_uploads"
REPORTS_DIR = Path(tempfile.gettempdir()) / "legaltech_reports"
DB_PATH     = Path(tempfile.gettempdir()) / "legaltech.db"
MAX_FILE_MB = 10

UPLOAD_DIR.mkdir(exist_ok=True)
REPORTS_DIR.mkdir(exist_ok=True)

# ── Системный промпт юриста РБ ────────────────────────────────────
SYSTEM_PROMPT = """Ты опытный юрист Республики Беларусь со специализацией в гражданском и корпоративном праве.

ПРАВИЛА:
- Всегда ссылайся на конкретные статьи НПА Республики Беларусь
- Используй только белорусское право (ГК РБ, ТК РБ, ХПК РБ и др.)
- Никогда не применяй российское право
- Каждый риск обоснуй ссылкой на статью

ФОРМАТ ОТВЕТА (строго JSON, без markdown, только чистый JSON объект):
{
  "contract_type": "тип договора",
  "overall_risk": "низкий",
  "summary": "краткое резюме 2-3 предложения",
  "risks": [
    {
      "level": "высокий",
      "clause": "пункт договора или цитата",
      "issue": "описание проблемы",
      "law_reference": "ст. 402 ГК РБ",
      "recommendation": "конкретная рекомендация"
    }
  ],
  "missing_clauses": ["список важных условий которых нет"],
  "needs_lawyer": true
}"""

# ── База данных (SQLite) ─────────────────────────────────────────

async def init_db():
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                created_at TEXT DEFAULT (datetime('now')),
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                contract_type TEXT,
                overall_risk TEXT,
                result_json TEXT,
                report_path TEXT,
                error_text TEXT,
                llm_provider TEXT,
                llm_model TEXT
            )
        """)
        await db.commit()
    logger.info(f"БД: {DB_PATH}")


async def db_create(doc_id: str, filename: str, file_path: str, provider: str, model: str):
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            "INSERT INTO documents (id, filename, file_path, status, llm_provider, llm_model) VALUES (?,?,?,'pending',?,?)",
            (doc_id, filename, str(file_path), provider, model)
        )
        await db.commit()


async def db_update(doc_id: str, **kwargs):
    if not kwargs:
        return
    sets = ", ".join(f"{k} = ?" for k in kwargs)
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(f"UPDATE documents SET {sets} WHERE id = ?", [*kwargs.values(), doc_id])
        await db.commit()


async def db_get(doc_id: str) -> Optional[dict]:
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)) as cur:
            row = await cur.fetchone()
            if not row:
                return None
            data = dict(row)
            if data.get("result_json"):
                data["result_json"] = json.loads(data["result_json"])
            return data

# ── Парсинг документов ────────────────────────────────────────────

def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        doc = fitz.open(str(file_path))
        text = "\n".join(p.get_text() for p in doc)
        doc.close()
    elif ext in (".docx", ".doc"):
        from docx import Document as DocxDoc
        doc = DocxDoc(str(file_path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts)
    else:
        raise ValueError(f"Формат не поддерживается: {ext}")
    if len(text.strip()) < 50:
        raise ValueError("Не удалось извлечь текст — файл пустой или защищён паролем")
    return text

# ── Генерация DOCX отчёта ─────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_report(result: dict, filename: str, provider: str, model: str) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    h = doc.add_heading(f"Анализ договора — {result['contract_type']}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Файл: {filename}")
    doc.add_paragraph(f"AI: {provider} / {model}")

    p = doc.add_paragraph()
    p.add_run("Общий уровень риска: ").bold = True
    run = p.add_run(result["overall_risk"].upper())
    run.bold = True
    run.font.size = Pt(14)
    clr = {"высокий": RGBColor(0xCC,0,0), "средний": RGBColor(0xFF,0x99,0), "низкий": RGBColor(0,0x80,0)}
    run.font.color.rgb = clr.get(result["overall_risk"].lower(), RGBColor(0,0,0))

    doc.add_paragraph()
    doc.add_heading("Резюме", level=2)
    doc.add_paragraph(result["summary"])
    doc.add_paragraph()

    doc.add_heading("Выявленные риски", level=2)
    risks = result.get("risks", [])
    if risks:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Уровень", "Пункт договора", "Проблема", "Норма РБ", "Рекомендация"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], "D9D9D9")
        bg = {"высокий": "FFCCCC", "средний": "FFF3CD", "низкий": "D4EDDA"}
        for risk in risks:
            row = tbl.add_row().cells
            row[0].text = risk.get("level", "")
            row[1].text = risk.get("clause", "")[:300]
            row[2].text = risk.get("issue", "")
            row[3].text = risk.get("law_reference", "")
            row[4].text = risk.get("recommendation", "")
            color = bg.get(risk.get("level", "").lower(), "FFFFFF")
            for cell in row:
                _set_cell_bg(cell, color)
    else:
        doc.add_paragraph("Существенных рисков не выявлено.")

    doc.add_paragraph()
    missing = result.get("missing_clauses", [])
    if missing:
        doc.add_heading("Отсутствующие условия", level=2)
        for clause in missing:
            doc.add_paragraph(clause, style="List Bullet")
        doc.add_paragraph()

    doc.add_heading("Вывод", level=2)
    p = doc.add_paragraph()
    p.add_run("Необходима консультация юриста: ").bold = True
    answer = "ДА" if result.get("needs_lawyer") else "НЕТ"
    run = p.add_run(answer)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC,0,0) if result.get("needs_lawyer") else RGBColor(0,0x80,0)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Подготовлено: LexAI.by | {provider} / {model}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66,0x66,0x66)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Пайплайн анализа ─────────────────────────────────────────────

async def process_contract(doc_id: str, file_path: Path, filename: str):
    llm = get_llm_provider()
    logger.info(f"[{doc_id}] Старт | {llm.name}/{llm.model} | {filename}")
    await db_update(doc_id, status="processing")

    try:
        # 1. Извлечь текст
        text = extract_text(file_path)
        logger.info(f"[{doc_id}] Текст: {len(text)} символов")

        # 2. Анализ через выбранный LLM провайдер
        user_msg = (
            f"Проанализируй следующий договор на соответствие законодательству "
            f"Республики Беларусь. Верни строго JSON без markdown.\n\n"
            f"ТЕКСТ ДОГОВОРА:\n{text[:12000]}"
        )
        response = await llm.chat(system=SYSTEM_PROMPT, user=user_msg)
        raw = response.text.strip()

        # Чистим от markdown и лишнего текста
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            raw = match.group()

        result = json.loads(raw)
        logger.info(f"[{doc_id}] Анализ готов: риск={result.get('overall_risk')}")

        # 3. Генерация DOCX
        report_bytes = generate_report(result, filename, llm.name, llm.model)
        report_path = REPORTS_DIR / f"{doc_id}.docx"
        report_path.write_bytes(report_bytes)

        await db_update(
            doc_id,
            status="done",
            contract_type=result.get("contract_type", "Неизвестный тип"),
            overall_risk=result.get("overall_risk", "средний"),
            result_json=json.dumps(result, ensure_ascii=False),
            report_path=str(report_path),
        )
        logger.info(f"[{doc_id}] Готово!")

    except Exception as e:
        logger.error(f"[{doc_id}] Ошибка: {e}")
        await db_update(doc_id, status="error", error_text=str(e))

# ── FastAPI ───────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    llm = get_llm_provider()
    emb = get_embedding_provider()
    llm_health = await llm.health()
    emb_health = await emb.health()
    logger.info(f"LLM: {llm_health}")
    logger.info(f"Embeddings: {emb_health}")
    if not llm_health.available:
        logger.warning(f"⚠️  LLM недоступен: {llm_health.error}")
    yield


app = FastAPI(
    title="LegalTech RB API",
    description="AI-анализ договоров по законодательству РБ. Провайдер: " + os.getenv("LLM_PROVIDER", "ollama"),
    version="1.1.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
async def health():
    llm = get_llm_provider()
    emb = get_embedding_provider()
    llm_h = await llm.health()
    emb_h = await emb.health()
    return {
        "status": "ok",
        "version": "1.1.0",
        "llm": {"provider": llm_h.name, "model": llm_h.model, "available": llm_h.available, "error": llm_h.error},
        "embeddings": {"provider": emb_h.name, "model": emb_h.model, "available": emb_h.available},
        "providers_available": ["ollama ✅", "groq 🔌 (free)", "anthropic 🔌", "openai 🔌"],
        "switch_provider": "Измените LLM_PROVIDER в .env: ollama | groq | anthropic | openai",
    }


@app.post("/api/upload")
async def upload(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(400, "Имя файла обязательно")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("pdf", "docx", "doc"):
        raise HTTPException(400, f"Неподдерживаемый формат '{ext}'. Нужен PDF или DOCX")

    data = await file.read()
    if len(data) > MAX_FILE_MB * 1024 * 1024:
        raise HTTPException(413, f"Файл > {MAX_FILE_MB} МБ")
    if len(data) < 100:
        raise HTTPException(400, "Файл пустой или повреждён")

    llm = get_llm_provider()
    doc_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{doc_id}.{ext}"
    file_path.write_bytes(data)

    await db_create(doc_id, file.filename, str(file_path), llm.name, llm.model)
    background_tasks.add_task(process_contract, doc_id, file_path, file.filename)

    logger.info(f"Загружен: {file.filename} → {doc_id} | провайдер: {llm.name}/{llm.model}")
    return {"task_id": doc_id, "status": "processing", "message": "Файл принят в обработку",
            "llm_provider": llm.name, "llm_model": llm.model}


@app.get("/api/task/{task_id}")
async def task_status(task_id: str):
    doc = await db_get(task_id)
    if not doc:
        raise HTTPException(404, "Задача не найдена")

    resp = {
        "task_id": task_id,
        "status": doc["status"],
        "filename": doc["filename"],
        "contract_type": doc.get("contract_type"),
        "overall_risk": doc.get("overall_risk"),
        "created_at": doc.get("created_at"),
        "llm_provider": doc.get("llm_provider"),
        "llm_model": doc.get("llm_model"),
    }
    if doc["status"] == "done":
        resp["result"] = doc.get("result_json")
        resp["download_url"] = f"/api/download/{task_id}"
    elif doc["status"] == "error":
        resp["error_message"] = doc.get("error_text")
    return resp


@app.get("/api/download/{task_id}")
async def download(task_id: str):
    doc = await db_get(task_id)
    if not doc:
        raise HTTPException(404, "Задача не найдена")
    if doc["status"] != "done":
        raise HTTPException(400, f"Отчёт не готов. Статус: {doc['status']}")

    report_path = Path(doc["report_path"])
    if not report_path.exists():
        raise HTTPException(404, "Файл отчёта не найден")

    original = doc["filename"].rsplit(".", 1)[0]
    return FileResponse(
        path=str(report_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"report_{original}.docx",
    )


if __name__ == "__main__":
    import uvicorn
    provider = os.getenv("LLM_PROVIDER", "ollama")
    logger.info(f"🚀 LegalTech RB | провайдер: {provider}")
    logger.info(f"   Сменить AI: LLM_PROVIDER=groq|anthropic|openai python main_local.py")
    uvicorn.run("main_local:app", host="0.0.0.0", port=8000, reload=False, log_level="warning")
